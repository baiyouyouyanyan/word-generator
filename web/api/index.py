import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import pandas as pd
import shutil
import uuid
from datetime import datetime
import json

def step1_replace_and_highlight(template_path: str, excel_path: str, output_path: str) -> int:
    """核心函数：处理模板中的 {占位符}，替换文字并只把替换部分标黄"""
    df_data = pd.read_excel(excel_path, sheet_name="data", header=0)
    data_dict = {
        str(k).strip(): str(v).strip()
        for k, v in zip(df_data.iloc[:, 0], df_data.iloc[:, 1])
        if pd.notna(k)
    }
    
    doc = Document(template_path)
    replaced = 0
    
    for para in doc.paragraphs:
        for key, value in data_dict.items():
            placeholder = f"{{{key}}}"
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        replaced += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_dict.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    replaced += 1
    
    doc.save(output_path)
    return replaced

UPLOAD_DIR = "/tmp/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def handler(environ, start_response):
    """Vercel Serverless Function Handler"""
    path = environ.get('PATH_INFO', '')
    method = environ.get('REQUEST_METHOD', 'GET')
    
    if path == '/api/generate' and method == 'POST':
        return handle_generate(environ, start_response)
    
    elif path.startswith('/api/download/'):
        return handle_download(environ, start_response, path)
    
    elif path == '/' or path == '/index.html':
        return serve_index(environ, start_response)
    
    else:
        return json_response(start_response, {"error": "Not found"}, 404)

def handle_generate(environ, start_response):
    """处理文件上传并生成报告"""
    try:
        content_length = int(environ.get('CONTENT_LENGTH', 0))
        content_type = environ.get('CONTENT_TYPE', '')
        
        boundary = None
        if 'multipart/form-data' in content_type:
            boundary = content_type.split('boundary=')[-1]
        
        if not boundary:
            return json_response(start_response, {"error": "需要 multipart/form-data"}, 400)
        
        post_data = environ['wsgi.input'].read(content_length)
        
        word_content = None
        excel_content = None
        
        parts = post_data.split(f'--{boundary}'.encode())
        for part in parts:
            if b'filename=' not in part:
                continue
            
            filename_match = part.decode('utf-8', errors='ignore').find('filename="')
            if filename_match == -1:
                continue
            
            header_end = part.find(b'\r\n\r\n')
            if header_end == -1:
                continue
            
            header = part[:header_end].decode('utf-8', errors='ignore')
            content = part[header_end + 4:]
            
            if 'template.docx' in header or 'word' in header.lower():
                word_content = content
            elif 'data.xlsx' in header or 'excel' in header.lower():
                excel_content = content
        
        if not word_content or not excel_content:
            return json_response(start_response, {"error": "缺少文件"}, 400)
        
        job_id = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
        job_dir = os.path.join(UPLOAD_DIR, job_id)
        os.makedirs(job_dir, exist_ok=True)
        
        word_path = os.path.join(job_dir, "template.docx")
        excel_path = os.path.join(job_dir, "data.xlsx")
        output_path = os.path.join(job_dir, "output.docx")
        
        with open(word_path, 'wb') as f:
            f.write(word_content)
        with open(excel_path, 'wb') as f:
            f.write(excel_content)
        
        try:
            replaced_count = step1_replace_and_highlight(word_path, excel_path, output_path)
            
            response_data = {
                "success": True,
                "file_url": f"/api/download/{job_id}/output.docx",
                "replaced_count": replaced_count
            }
            
            return json_response(start_response, response_data, 200)
            
        except Exception as e:
            return json_response(start_response, {"error": str(e)}, 500)
        finally:
            try:
                os.remove(word_path)
                os.remove(excel_path)
            except:
                pass
    
    except Exception as e:
        return json_response(start_response, {"error": str(e)}, 500)

def handle_download(environ, start_response, path):
    """下载生成的文件"""
    parts = path.split('/api/download/')
    if len(parts) < 2:
        return json_response(start_response, {"error": "Invalid path"}, 400)
    
    job_id = parts[1].split('/')[0]
    filename = parts[1].split('/')[1] if '/' in parts[1] else 'output.docx'
    
    file_path = os.path.join(UPLOAD_DIR, job_id, filename)
    
    if not os.path.exists(file_path):
        return json_response(start_response, {"error": "File not found"}, 404)
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
        
        shutil.rmtree(os.path.join(UPLOAD_DIR, job_id), ignore_errors=True)
        
        status = '200 OK'
        response_headers = [
            ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            ('Content-Disposition', f'attachment; filename="{filename}"'),
            ('Content-Length', str(len(content)))
        ]
        
        start_response(status, response_headers)
        return [content]
    except Exception as e:
        return json_response(start_response, {"error": str(e)}, 500)

def serve_index(environ, start_response):
    """返回HTML页面"""
    html_path = os.path.join(os.path.dirname(__file__), 'index.html')
    
    if os.path.exists(html_path):
        with open(html_path, 'r', encoding='utf-8') as f:
            content = f.read()
    else:
        content = "<html><body><h1>请将 index.html 放在项目根目录</h1></body></html>"
    
    status = '200 OK'
    response_headers = [
        ('Content-Type', 'text/html; charset=utf-8'),
        ('Content-Length', str(len(content.encode('utf-8'))))
    ]
    start_response(status, response_headers)
    return [content.encode('utf-8')]

def json_response(start_response, data, status_code=200):
    """返回JSON响应"""
    json_str = json.dumps(data, ensure_ascii=False)
    json_bytes = json_str.encode('utf-8')
    
    status = f'{status_code} OK'
    response_headers = [
        ('Content-Type', 'application/json; charset=utf-8'),
        ('Content-Length', str(len(json_bytes)))
    ]
    start_response(status, response_headers)
    return [json_bytes]
