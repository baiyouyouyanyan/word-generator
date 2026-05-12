from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from openpyxl import load_workbook
import os
import uuid
import shutil
from datetime import datetime

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def step1_replace_and_highlight(template_path: str, excel_path: str, output_path: str) -> int:
    wb = load_workbook(excel_path, data_only=True)
    ws = wb['data']
    
    data_dict = {}
    for row in ws.iter_rows(min_row=2):
        key = str(row[0].value).strip() if row[0].value else ''
        value = str(row[1].value).strip() if row[1].value else ''
        if key:
            data_dict[key] = value
    
    doc = Document(template_path)
    replaced = 0
    
    for para in doc.paragraphs:
        for key, value in data_dict.items():
            placeholder = f"{{{key}}}"
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        replaced += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_dict.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    replaced += 1
    
    doc.save(output_path)
    return replaced

@app.get("/")
async def root():
    return FileResponse("index.html")

@app.post("/api/generate")
async def generate_report(word: UploadFile = File(...), excel: UploadFile = File(...)):
    if not word.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="请上传Word文档(.docx)")
    
    if not excel.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="请上传Excel文件(.xlsx)")
    
    job_id = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
    job_dir = os.path.join(UPLOAD_DIR, job_id)
    os.makedirs(job_dir, exist_ok=True)
    
    word_path = os.path.join(job_dir, "template.docx")
    excel_path = os.path.join(job_dir, "data.xlsx")
    output_path = os.path.join(job_dir, "output.docx")
    
    with open(word_path, "wb") as f:
        shutil.copyfileobj(word.file, f)
    
    with open(excel_path, "wb") as f:
        shutil.copyfileobj(excel.file, f)
    
    try:
        replaced_count = step1_replace_and_highlight(word_path, excel_path, output_path)
        
        return {
            "success": True,
            "file_url": f"/api/download/{job_id}/output.docx",
            "replaced_count": replaced_count
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")
    finally:
        try:
            os.remove(word_path)
            os.remove(excel_path)
        except:
            pass

@app.get("/api/download/{job_id}/{filename}")
async def download_file(job_id: str, filename: str):
    file_path = os.path.join(UPLOAD_DIR, job_id, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
